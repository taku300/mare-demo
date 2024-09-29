import streamlit as st
import os
from langchain.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, AIMessage, SystemMessage
from docx import Document
from io import BytesIO
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from PIL import Image
import MeCab  # MeCabを使用して日本語形態素解析を行う
from collections import Counter
from configs import const

# OpenAI APIキーを環境変数から設定
os.environ["OPENAI_API_KEY"] = const.OPENAI_API_KEY

# ChatOpenAIのインスタンスを作成
llm = ChatOpenAI(temperature=0, model_name="gpt-4")

# 日本語フォントのパスを指定する
font_path = '/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc'

# システムプロンプトを定義
system_prompt = """
# 役割
あなたは顧客にニーズを見つけ出すスペシャリストです。
お客様の心の中にあるインサイトを見つけ出すように質問の深掘りをしてください。

# 振る舞い
- 丁寧な振る舞いをする
- 質問は端的に返す。
- 十分深掘りが完了と判断したら「ご回答ありがとうございました。」と返答する。
"""

# インサイト分析プロンプトを定義
insight_analysis_prompt = """
# 役割
あなたは顧客のインサイト分析のスペシャリストです。
以下の対話履歴を元に、顧客の潜在的なニーズやインサイトを分析してください。
レポートは300字以内でまとめてください。

# 対話履歴
{conversation}

# インサイト分析結果
"""

# 新商品のアイデアプロンプトを定義
new_product_idea_prompt = """
# 役割
あなたは商品開発のスペシャリストです。
以下の対話履歴と顧客インサイトを元に、今後の新商品のアイデアを提案してください。
アイデアは具体的でかつ顧客の潜在ニーズに沿ったものである必要があります。

# 対話履歴
{conversation}

# 顧客インサイト
{insights}

# 新商品のアイデア
"""

# 形態素解析を行い、名詞、動詞、形容詞、形容動詞、副詞の原型を抽出する関数
def extract_keywords(text):
    mecab = MeCab.Tagger("-Ochasen")
    parsed = mecab.parse(text)

    # 名詞、動詞、形容詞、形容動詞、副詞を抽出し、基本形（原型）を取得する
    keywords = []
    for line in parsed.splitlines():
        if line == 'EOS':
            break
        parts = line.split("\t")
        if len(parts) > 3:
            base = parts[2]  # 基本形（原型）を取得
            pos = parts[3]   # 品詞情報を取得
            # 名詞、動詞、形容詞、形容動詞、副詞のみを対象にする
            if "名詞" in pos or "動詞" in pos or "形容詞" in pos or "形容動詞" in pos or "副詞" in pos:
                keywords.append(base)  # 基本形を使用
    return keywords

# WordCloudを生成する関数
def generate_wordcloud(text):
    # 形態素解析で単語を抽出
    words = extract_keywords(text)

    # 不要な単語を除外
    stopwords = {'user', 'assistant', 'です', 'ます', 'する', 'ある'}  # 必要に応じて増やす
    words = [word for word in words if word not in stopwords]

    # 単語の出現頻度をカウント
    word_freq = Counter(words)

    # WordCloudを生成
    wordcloud = WordCloud(width=400, height=300, background_color='white', font_path=font_path).generate_from_frequencies(word_freq)

    # WordCloud画像を保存する
    image_io = BytesIO()
    wordcloud.to_image().save(image_io, format='PNG')
    image_io.seek(0)

    return image_io

def analyze_insights(conversation):
    prompt = insight_analysis_prompt.format(conversation=conversation)
    analysis_response = llm([SystemMessage(content=system_prompt), HumanMessage(content=prompt)])
    return analysis_response.content

def generate_new_product_ideas(conversation, insights):
    prompt = new_product_idea_prompt.format(conversation=conversation, insights=insights)
    idea_response = llm([SystemMessage(content=system_prompt), HumanMessage(content=prompt)])
    return idea_response.content

# ストリーミング再生に対応した応答生成関数
def generate_response(messages):
    for chunk in llm.stream(messages):
        yield chunk.content  # ストリーミング応答を逐次的に返す

def save_conversation_to_word(messages):
    doc = Document()
    doc.add_heading('インサイト分析レポート', 0)

    # 対話履歴を作成
    doc.add_heading('対話履歴', level=1)
    conversation = ""
    for message in messages:
        if message["role"] == "user":
            doc.add_heading('User:', level=2)
            doc.add_paragraph(message["content"])
            conversation += f"{message['content']}\n"  # Userのメッセージのみを追加
        elif message["role"] == "assistant":
            doc.add_heading('Assistant:', level=2)
            doc.add_paragraph(message["content"])

    # 顧客インサイトを作成
    insights = analyze_insights(conversation)
    doc.add_heading('顧客インサイト', level=1)
    doc.add_paragraph(insights)

    # 新商品のアイデアを生成
    new_product_ideas = generate_new_product_ideas(conversation, insights)
    doc.add_heading('今後の新商品のアイデア', level=1)
    doc.add_paragraph(new_product_ideas)

    # Word Cloudの生成と追加 (Userからの回答のみを使う)
    doc.add_heading('よく使われた単語（Word Cloud）', level=1)

    # Userのメッセージのみを対象にWord Cloudを生成
    user_conversation = "\n".join([msg["content"] for msg in messages if msg["role"] == "user"])
    wordcloud_image = generate_wordcloud(user_conversation)

    # Wordファイルに画像としてWordCloudを追加
    image = Image.open(wordcloud_image)
    image_path = '/tmp/wordcloud.png'
    image.save(image_path)  # 一時的に画像ファイルを保存

    doc.add_picture(image_path)

    # メモリに保存
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    return byte_io

def show():
    st.title("アンケート分析AI")
    st.markdown(("### 新商品のカレーの味はどうでしたか？"))

    # セッションステートにメッセージを保存
    if "messages" not in st.session_state:
        st.session_state.messages = []
        # 初期化時にシステムプロンプトをメッセージリストに追加
        st.session_state.messages.append({"role": "system", "content": system_prompt})

    # メッセージの表示（システムプロンプトは表示しない）
    for message in st.session_state.messages:
        if message["role"] != "system":
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

    # ユーザーの入力を受け取る
    prompt = st.chat_input("文章を入力してください。")
    if prompt:
        with st.chat_message("user"):
            st.markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})

        # メッセージリストにシステムプロンプトを含める
        messages = [SystemMessage(content=system_prompt)] + [
            HumanMessage(content=msg["content"]) for msg in st.session_state.messages if msg["role"] != "system"
        ]

        # ストリーミング応答をリアルタイムで表示
        with st.chat_message("assistant"):
            response_stream = generate_response(messages)
            # Streamlit の write_stream を使ってストリーミング再生
            full_response = st.write_stream(response_stream)

        # 最終応答をセッションに保存
        st.session_state.messages.append({"role": "assistant", "content": full_response})

    # Wordに保存するボタンを表示
    if st.button("レポートを作成"):
        word_data = save_conversation_to_word(st.session_state.messages)
        st.success("レポートを作成しました。")

        # ダウンロードボタンを追加
        st.download_button(
            label="レポートをダウンロード",
            data=word_data,
            file_name="report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    show()

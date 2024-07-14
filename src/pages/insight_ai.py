import streamlit as st
import os
from langchain.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, AIMessage, SystemMessage
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from configs import const

# OpenAI APIキーを環境変数から設定
os.environ["OPENAI_API_KEY"] = const.OPENAI_API_KEY

# ChatOpenAIのインスタンスを作成
llm = ChatOpenAI(temperature=0, model_name="gpt-4o")

# システムプロンプトを定義
system_prompt = """
# 役割
あなたは顧客にニーズを見つけ出すスペシャリストです。
お客様の心の中にあるインサイトを見つけ出すように質問の深掘りをしてください。

# 振る舞い
- 丁寧な振る舞いをする
- 質問は端的に返す。
- 十分深掘りが完了と判断したらしたら「ご回答ありがとうございました。」と返答する。
"""

# インサイト分析プロンプトを定義
insight_analysis_prompt = """
# 役割
あなたは顧客のインサイト分析のスペシャリストです。
以下の対話履歴を元に、顧客の潜在的なニーズやインサイトを分析してください。
出力の際はマークダウンは使わないでください。
レポートは300字以内でまとめてください。

# 対話履歴
{conversation}

# インサイト分析結果
"""

def generate_response(messages):
    response = llm(messages)
    return response.content

def analyze_insights(conversation):
    prompt = insight_analysis_prompt.format(conversation=conversation)
    analysis_response = llm([SystemMessage(content=system_prompt), HumanMessage(content=prompt)])
    return analysis_response.content

def save_conversation_to_word(messages, output_path):
    doc = Document()
    doc.add_heading('インサイト分析レポート', 0)

    conversation = ""
    for message in messages:
        if message["role"] == "user":
            doc.add_heading('User:', level=1)
        elif message["role"] == "assistant":
            doc.add_heading('Assistant:', level=1)
        else:
            continue  # Skip system messages

        doc.add_paragraph(message["content"])
        conversation += f"{message['role']}: {message['content']}\n"

    # インサイト分析を実行
    insights = analyze_insights(conversation)

    doc.add_heading('顧客インサイト', level=1)
    
    # インサイト分析結果をパースし、適切なスタイルを適用
    paragraphs = insights.split('\n')
    for paragraph in paragraphs:
        if paragraph.startswith("1.") or paragraph.startswith("2.") or paragraph.startswith("3."):
            p = doc.add_paragraph()
            run = p.add_run(paragraph)
            run.bold = True
        elif paragraph.startswith("**") and paragraph.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(paragraph[2:-2])
            run.bold = True
        else:
            doc.add_paragraph(paragraph)

    # 完全なファイルパスを作成
    file_path = os.path.join(output_path, "report.docx")
    doc.save(file_path)

def show():
    st.title("Insight AI")

    # セッションステートにメッセージを保存
    if "messages" not in st.session_state:
        st.session_state.messages = []
        # 初期化時にシステムプロンプトをメッセージリストに追加
        st.session_state.messages.append({"role": "system", "content": system_prompt})

    # メッセージの表示（システムプロンプトは表示しない）
    for message in st.session_state.messages:
        if message["role"] != "system":
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

    # ユーザーの入力を受け取る
    prompt = st.chat_input("文章を入力してください。")
    if prompt:
        with st.chat_message("user"):
            st.markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})

        # メッセージリストにシステムプロンプトを含める
        messages = [SystemMessage(content=system_prompt)] + [
            HumanMessage(content=msg["content"]) for msg in st.session_state.messages if msg["role"] != "system"
        ]

        response = generate_response(messages)
        
        with st.chat_message("assistant"):
            st.markdown(response)
        st.session_state.messages.append({"role": "assistant", "content": response})

    # Wordに保存するボタンを常に表示
    if st.button("レポートを作成"):
        save_conversation_to_word(st.session_state.messages, const.OUTPUT_PATH)
        st.success("レポートを作成しました。")

if __name__ == "__main__":
    show()
